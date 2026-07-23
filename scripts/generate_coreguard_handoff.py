#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install it with: pip install python-docx"
    ) from exc


def write_docx_from_markdown(md_text: str, out_path: Path) -> None:
    document = Document()
    lines = md_text.splitlines()
    in_code = False

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            p = document.add_paragraph()
            p.add_run(line).font.name = "Courier New"
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            continue

        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            document.add_paragraph(text, style="List Bullet")
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            document.add_paragraph(text, style="List Number")
            continue

        if "|" in line and not line.strip().startswith("|---"):
            document.add_paragraph(line.replace("|", "  |  "))
            continue

        document.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def create_handoff_zip(zip_path: Path, files: list[Path]) -> None:
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_path in files:
            if file_path.exists():
                zf.write(file_path, arcname=file_path.name)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate CoreGuard SRD handoff artifacts (.md, .docx, .zip)."
    )
    parser.add_argument(
        "--source-md",
        type=Path,
        default=Path("docs/CoreGuard_Elite_Copilot_SRD.md"),
        help="Source markdown SRD path.",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=Path("build/handoff"),
        help="Output directory for generated artifacts.",
    )
    parser.add_argument(
        "--logo",
        type=Path,
        default=None,
        help="Optional logo file to include in handoff zip.",
    )
    args = parser.parse_args()

    source_md = args.source_md
    if not source_md.exists():
        raise SystemExit(f"Source markdown not found: {source_md}")

    out_dir = args.out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    md_out = out_dir / "CoreGuard_Elite_Copilot_SRD.md"
    docx_out = out_dir / "CoreGuard_Elite_Copilot_SRD.docx"
    zip_out = out_dir / "CoreGuard_Elite_Copilot_Handoff.zip"

    markdown_text = source_md.read_text(encoding="utf-8")
    md_out.write_text(markdown_text, encoding="utf-8")
    write_docx_from_markdown(markdown_text, docx_out)

    handoff_files = [md_out, docx_out]
    if args.logo and args.logo.exists():
        handoff_files.append(args.logo)
    create_handoff_zip(zip_out, handoff_files)

    print(f"Generated: {md_out}")
    print(f"Generated: {docx_out}")
    print(f"Generated: {zip_out}")
    if args.logo and args.logo.exists():
        print(f"Included logo: {args.logo}")
    else:
        print("Logo not found; zip created without logo.")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Generate CoreGuard Elite SRD handoff artifacts (.md, .docx, .zip).

Compatible with:
  - Jupyter / Google Colab  (set OUT_DIR env var before running)
  - Local Python 3.10+      (pass --out-dir on the CLI)

Dependencies::

    pip install "python-docx>=0.8.11"

Usage examples::

    # Local — output to build/handoff/
    python tools/generate_coreguard_elite_handoff.py --out-dir build/handoff

    # With an optional logo on the title page
    python tools/generate_coreguard_elite_handoff.py \\
        --out-dir build/handoff \\
        --logo assets/coreguard_logo.png

    # Colab / Jupyter — just set OUT_DIR before running all cells
    # import os; os.environ["OUT_DIR"] = "/mnt/data"
"""
from __future__ import annotations

import argparse
import os
import re
import zipfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as _exc:
    raise SystemExit(
        "python-docx >= 0.8.11 is required.  "
        "Install it with:  pip install python-docx"
    ) from _exc

# ---------------------------------------------------------------------------
# SRD content
# Raw string — backslashes are literal, no escape expansion (intentional).
# Edit this constant to update the document without touching generation logic.
# ---------------------------------------------------------------------------

_CONTENT = r"""# CoreGuard Elite — Copilot Implementation SRD

**Document type:** Software Requirements Document and implementation handoff
**Audience:** GitHub Copilot, Cursor, Gemini, Android engineers, QA engineers, security reviewers
**Product owner:** CoreGuard Technologies
**Application:** CoreGuard Elite
**AI analyst:** Quilla Intelligence
**Security scanner:** Nemesis Scanner
**Canonical package:** com.app.coldboarcoreguard
**Document status:** Implementation directive; preserve working features and extend incrementally
**Version:** 1.0
**Date:** 2026-07-23

---

## 1. Purpose

This SRD gives an implementation AI or engineer enough detail to continue CoreGuard Elite without
discarding, renaming, weakening, or inaccurately representing the work already completed.

The current project is a functional Android security-intelligence MVP source foundation. It
includes a multi-stage evidence pipeline, Quilla risk correlation, Room-backed scan history,
remediation status, VPN and kill-switch scaffolding, StrongBox/Keystore support, RASP signals,
enterprise package containment, signed IOC foundations, branding assets, and a Compose UI.

The next implementation must:

1. Preserve all completed security logic.
2. Compile and test the application before adding broad new features.
3. Refactor safely rather than rewrite the project.
4. Distinguish implemented behavior from planned behavior.
5. Never market heuristics as proof of malware or compromise.
6. Maintain the established CoreGuard visual and language system.
7. Keep all potentially disruptive actions explicit, reversible, and user controlled.

---

## 2. Product hierarchy and naming

| Name | Meaning | Usage |
|:---|:---|:---|
| CoreGuard Technologies | Company and publisher | Legal, website, Play Console, business documents |
| CoreGuard Elite | Android application | Launcher label, store listing, in-app product name |
| Quilla Intelligence | Calm AI security analyst and evidence-correlation layer | Explanations, confidence, summaries, recommendations |
| Nemesis Scanner | Multi-stage device and application verification workflow | Scan screen, scan history, findings |
| Network Shield | VPN-based network protection feature family | VPN UI and future packet/DNS protections |
| Integrity Index | Explainable 0-100 security consistency score | Dashboard, reports, scan sessions |
| Integrity Timeline | Chronological security and response history | Timeline and investigation features |
| Observations | Evidence-backed findings; preferred over alerts | Results, details, reports |

Do not rename Quilla, Nemesis, Integrity Index, Integrity Timeline, Network Shield, or CoreGuard
Elite without product-owner approval.

---

## 3. Product vision

CoreGuard Elite is a premium Android security-intelligence application designed to give users
clarity and peace of mind.

It should feel like:

> A scientific instrument, a modern intelligence system, and a protective observatory.

It must not feel like:

- a fear-based antivirus advertisement;
- a fake one-tap hacker detector;
- an occult or supernatural product;
- an app that promises impossible Android capabilities;
- a cluttered technical console that ordinary users cannot understand.

Core philosophy:

- Observe before concluding.
- Verify before trusting.
- Explain before recommending.
- Respond proportionally.
- Confidence comes from evidence.

Core operational model: Observe → Correlate → Explain → Respond → Verify
User-facing wording: Observe → Interpret → Protect → Verify

---

## 4. Current implementation status

### 4.1 Implemented now

- Jetpack Compose Android application
- Android namespace and application ID: com.app.coldboarcoreguard
- Minimum Android SDK 26
- Target and compile SDK 35
- Java/Kotlin 17
- Material 3
- Navigation Compose
- Coroutines
- Room database with KSP
- Billing client scaffold
- Native C++ RASP component through CMake
- Multi-stage Nemesis scan orchestration
- Cancellable scan progress with eight named stages
- Granted-sensitive-permission analysis
- Installed-package analysis subject to Android package visibility
- Signing-certificate SHA-256 collection and analysis
- APK SHA-256 collection and signed IOC matching
- Installation-source analysis
- Accessibility service and Device Administrator observations
- Root, test-key, debugger, emulator, and related integrity signals
- Cryptographically verified IOC feed foundation
- Quilla correlation and Integrity Index assessment
- Persistent baseline comparison
- Room-backed scan sessions
- Room-backed observations and evidence
- Searchable scan history
- Severity filters
- Remediation status tracking
- Persistent timeline events
- Report-building foundation
- VPN service scaffold
- Explicit emergency kill switch with timeout and recovery safeguards
- StrongBox AES key request with Android Keystore fallback
- Key self-test and key rotation foundations
- Native and Kotlin RASP signals
- Device Owner package suspension with protected-package guardrails
- Security self-test
- Branding images and in-app emblem
- Electric teal/cyan, pewter/turquoise, and restrained-gold theme direction
- Gold particle background effect
- GitHub Actions Android CI workflow
- Unit tests for Quilla, remediation, severity, scan stages, and remediation statuses

### 4.2 Not yet verified or complete

The following must not be represented as production-complete:

- A fully successful clean Gradle build on the final branch
- Full instrumentation and physical-device testing
- A real traffic-forwarding VPN
- DNS filtering or malicious-domain blocking
- Deep packet inspection
- Universal malware detection
- Pegasus detection
- Forensic acquisition
- Production signed IOC backend and production key
- Play Integrity backend verification
- Production Firebase configuration
- Server-validated billing entitlements
- Signed release AAB
- Google Play policy approval
- Final privacy policy and Data Safety declarations
- Final screenshots captured from a running build

---

## 5. Required technology baseline

### 5.1 Current build configuration

```
namespace         = com.app.coldboarcoreguard
applicationId     = com.app.coldboarcoreguard
minSdk            = 26
compileSdk        = 35
targetSdk         = 35
versionCode       = 1
versionName       = 1.0.0
jvmTarget         = 17
```

### 5.2 Migration note

This SRD targets com.app.coldboarcoreguard (the Elite package). If continuing work on the
com.coldboar.coreguard baseline, migration must be a controlled plan with explicit test and
policy validation gates before any package or SDK change is merged.

---

## 6. Implementation constraints

1. **Preserve completed security logic.** Do not remove, rename, or weaken any evaluator,
   scanner, VPN component, RASP signal, or Keystore integration without an explicit migration plan.
2. **Compile and test before broad new feature additions.** Run the full Gradle test suite
   and confirm zero failures before starting new feature work.
3. **Refactor safely rather than rewrite.** Prefer incremental, targeted changes.
4. **Distinguish implemented from planned.** Mark anything not yet validated with a clear
   status indicator in comments and documentation.
5. **Never market heuristics as proof.** Quilla assessments are confidence-weighted evidence
   summaries, not guarantees of compromise or safety.
6. **Keep disruptive actions reversible.** Kill switch, package suspension, and key rotation
   must all be user-initiated, time-bounded, and recoverable.
7. **Do not rename protected product names** without product-owner approval.
"""

# ---------------------------------------------------------------------------
# Metadata for the title-page table (kept separate from _CONTENT so the
# title page and the body each have exactly one copy of this information).
# ---------------------------------------------------------------------------

_METADATA: list[tuple[str, str]] = [
    ("Document type", "Software Requirements Document and implementation handoff"),
    ("Audience", "GitHub Copilot, Cursor, Gemini, Android engineers, QA, security reviewers"),
    ("Product owner", "CoreGuard Technologies"),
    ("Application", "CoreGuard Elite"),
    ("AI analyst", "Quilla Intelligence"),
    ("Security scanner", "Nemesis Scanner"),
    ("Canonical package", "com.app.coldboarcoreguard"),
    ("Document status", "Implementation directive; preserve working features and extend incrementally"),
    ("Version", "1.0"),
    ("Date", "2026-07-23"),
]

# ---------------------------------------------------------------------------
# Document-building helpers
# ---------------------------------------------------------------------------


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a thin grey bottom border on an empty paragraph (visual HR)."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")       # border thickness (half-points)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_inline_bold(para, text: str) -> None:
    """Append runs to *para* applying bold to **text** spans inline."""
    # Split on double-asterisk markers; odd-indexed segments are bold.
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if part:  # skip empty strings from consecutive markers
            run = para.add_run(part)
            run.bold = (i % 2 == 1)


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    """Convert raw markdown table lines into a 2-D list of cell strings.

    Separator rows (e.g. |:---|:---|) are excluded so the result contains
    only header and data rows.
    """
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        # Skip separator rows: they contain only dashes, colons, pipes, spaces
        if re.match(r"^\|[-:| ]+\|$", stripped):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def _flush_md_table(doc: Document, table_lines: list[str]) -> None:
    """Convert buffered markdown table lines into a styled Word table."""
    rows = _parse_md_table(table_lines)
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(col_count):
            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = cell_text
            # Bold the header row (first row of the table)
            if r_idx == 0:
                for cell_para in cell.paragraphs:
                    for run in cell_para.runs:
                        run.bold = True


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------


def build_document(logo_path: Path) -> Document:
    """Build and return a styled Word document from _CONTENT.

    The title page (logo + title + metadata table) is generated first.
    Body content is parsed from _CONTENT starting at the first ## heading,
    so the metadata block at the top of _CONTENT is not duplicated.
    """
    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────

    if logo_path.exists():
        # Embed logo centered above the title
        doc.add_picture(str(logo_path), width=Inches(2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Logo is cosmetic — continue without it
        print(f"Warning: logo not found at {logo_path!s} — skipping.")

    title_para = doc.add_heading("CoreGuard Elite", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_para = doc.add_heading("Copilot Implementation SRD", level=1)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # Metadata table — bold keys in the left column
    meta_tbl = doc.add_table(rows=len(_METADATA), cols=2)
    meta_tbl.style = "Table Grid"
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(_METADATA):
        row = meta_tbl.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        for cell_para in row.cells[0].paragraphs:
            for run in cell_para.runs:
                run.bold = True

    doc.add_page_break()

    # ── Body: parse _CONTENT from the first ## heading onward ─────────────────
    # This avoids duplicating the document-type / audience header block that
    # is already present in the title-page metadata table above.

    content_lines = _CONTENT.splitlines()
    # Find the index of the first section heading (## …)
    body_start = next(
        (i for i, ln in enumerate(content_lines) if ln.strip().startswith("## ")),
        0,  # fall back to the very start if no ## heading is found
    )

    # Parsing state
    in_code_block = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        """Flush any accumulated markdown table lines to the document."""
        if table_lines:
            _flush_md_table(doc, list(table_lines))
            table_lines.clear()

    for line in content_lines[body_start:]:
        stripped = line.strip()

        # ── Code fences (``` … ```) ──────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code_block:
                # Closing fence — write buffered code as monospaced paragraph
                if code_lines:
                    para = doc.add_paragraph("\n".join(code_lines))
                    para.style = "No Spacing"
                    for run in para.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                code_lines.clear()
                in_code_block = False
            else:
                # Opening fence — flush any pending table first
                flush_table()
                in_code_block = True
            continue  # fence marker itself is not rendered

        if in_code_block:
            # Preserve indentation inside fenced blocks
            code_lines.append(line)
            continue

        # ── Markdown tables (lines starting with |) ──────────────────────────
        if stripped.startswith("|"):
            table_lines.append(line)
            continue
        else:
            # Non-table line — commit any pending table before proceeding
            flush_table()

        # ── Standard markdown elements ────────────────────────────────────────

        if not stripped:
            doc.add_paragraph()

        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)

        elif stripped.startswith("# "):
            # Top-level heading already on title page — render as level-1 in body
            doc.add_heading(stripped[2:], level=1)

        elif stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(para, stripped[2:])

        elif re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_inline_bold(para, text)

        elif stripped.startswith("> "):
            para = doc.add_paragraph(stripped[2:])
            try:
                # "Intense Quote" exists in the default Word template
                para.style = "Intense Quote"
            except KeyError:
                # Graceful fallback if the style is absent from the base template
                para.style = "Normal"

        elif stripped == "---":
            _add_horizontal_rule(doc)

        else:
            # Default: plain paragraph with inline bold support
            para = doc.add_paragraph()
            _add_inline_bold(para, stripped)

    # Flush any trailing table or unclosed code block
    flush_table()
    if in_code_block and code_lines:
        para = doc.add_paragraph("\n".join(code_lines))
        para.style = "No Spacing"
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    return doc


# ---------------------------------------------------------------------------
# I/O helpers
# ---------------------------------------------------------------------------


def save_md(path: Path) -> None:
    """Write the SRD markdown content to *path*."""
    path.write_text(_CONTENT, encoding="utf-8")
    print(f"Saved: {path}")


def save_docx(doc: Document, path: Path) -> None:
    """Persist the Word document to *path*."""
    doc.save(str(path))
    print(f"Saved: {path}")


def build_zip(zip_path: Path, *artifact_paths: Path) -> None:
    """Bundle *artifact_paths* into a ZIP archive at *zip_path*.

    Missing files produce a warning rather than an exception so that a
    partial handoff can still be created (e.g. when the logo is absent).
    """
    existing = [f for f in artifact_paths if f.exists()]
    missing = [f for f in artifact_paths if not f.exists()]
    if missing:
        print(f"Warning: skipping missing files: {[str(f) for f in missing]}")

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in existing:
            # Store only the filename — no directory structure inside the zip
            zf.write(f, arcname=f.name)

    print(f"Saved: {zip_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def _resolve_out_dir(cli_value: str | None) -> Path:
    """Resolve output directory: CLI arg > OUT_DIR env var > /mnt/data default."""
    if cli_value:
        return Path(cli_value)
    return Path(os.environ.get("OUT_DIR", "/mnt/data"))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate CoreGuard Elite SRD handoff artifacts (.md, .docx, .zip)."
    )
    parser.add_argument(
        "--out-dir",
        default=None,
        metavar="DIR",
        help=(
            "Output directory for generated artifacts. "
            "Overrides the OUT_DIR environment variable (default: /mnt/data)."
        ),
    )
    parser.add_argument(
        "--logo",
        default=None,
        metavar="PNG",
        help=(
            "Path to a logo PNG to embed on the .docx title page. "
            "Defaults to coreguard_logo_with_sleek_gold_accents.png in --out-dir."
        ),
    )
    args = parser.parse_args()

    out_dir = _resolve_out_dir(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    md_path = out_dir / "CoreGuard_Elite_Copilot_SRD.md"
    docx_path = out_dir / "CoreGuard_Elite_Copilot_SRD.docx"
    zip_path = out_dir / "CoreGuard_Elite_Copilot_Handoff.zip"

    # Prefer explicit --logo; fall back to the conventional filename beside out_dir
    logo_path = (
        Path(args.logo)
        if args.logo
        else out_dir / "coreguard_logo_with_sleek_gold_accents.png"
    )

    save_md(md_path)
    doc = build_document(logo_path)
    save_docx(doc, docx_path)
    build_zip(zip_path, md_path, docx_path, logo_path)
    print("Done — handoff package ready.")


if __name__ == "__main__":
    main()

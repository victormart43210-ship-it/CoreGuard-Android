#!/usr/bin/env python3
"""Generate an up-to-date CoreGuard SRD by introspecting the repository.

The script walks the repository to collect:
  - Build configuration from app/build.gradle.kts
  - Implemented Kotlin source components (main + mvt packages)
  - Unit-test coverage by file
  - Existing documentation files under docs/

It then writes a structured SRD markdown to the output path.
Optionally it converts the markdown to a .docx file and bundles both
into a .zip archive — the same output format used by
``scripts/generate_coreguard_handoff.py``.

Usage examples::

    # Markdown only (default output: build/srd/CoreGuard_SRD.md)
    python tools/generate_coreguard_srd.py

    # Markdown + docx + zip
    python tools/generate_coreguard_srd.py --docx --zip

    # Custom output directory and repo root
    python tools/generate_coreguard_srd.py --repo-root /path/to/repo --out-dir /tmp/srd
"""
from __future__ import annotations

import argparse
import datetime
import re
import zipfile
from pathlib import Path
from typing import NamedTuple

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


class BuildConfig(NamedTuple):
    namespace: str
    application_id: str
    min_sdk: str
    compile_sdk: str
    target_sdk: str
    version_code: str
    version_name: str
    jvm_target: str
    minify_release: bool


class ComponentInfo(NamedTuple):
    relative_path: str  # relative to repo root
    class_name: str
    package: str        # "main" or "mvt" or "ui"
    has_tests: bool


# ---------------------------------------------------------------------------
# Repository introspection helpers
# ---------------------------------------------------------------------------

_GRADLE_KV = re.compile(r'(\w+)\s*=\s*"?([^"\n]+)"?')
_MINIFY = re.compile(r"isMinifyEnabled\s*=\s*(true|false)")


def _parse_build_config(gradle_path: Path) -> BuildConfig:
    text = gradle_path.read_text(encoding="utf-8")

    def _get(key: str, default: str = "?") -> str:
        m = re.search(rf'{key}\s*=\s*"?([^"\n,)]+)"?', text)
        return m.group(1).strip() if m else default

    # isMinifyEnabled appears twice; we want the release block value
    minify_match = _MINIFY.findall(text)
    minify_release = minify_match[0] == "true" if minify_match else False

    return BuildConfig(
        namespace=_get("namespace"),
        application_id=_get("applicationId"),
        min_sdk=_get("minSdk"),
        compile_sdk=_get("compileSdk"),
        target_sdk=_get("targetSdk"),
        version_code=_get("versionCode"),
        version_name=_get("versionName"),
        jvm_target=_get("jvmTarget"),
        minify_release=minify_release,
    )


def _class_name_from_file(kt_file: Path) -> str:
    """Return the primary public class/object name declared in a Kotlin file.

    Only matches actual declarations at the start of a line (with optional
    visibility/modifier keywords), ignoring occurrences inside comments or
    KDoc blocks.
    """
    text = kt_file.read_text(encoding="utf-8", errors="replace")
    # Match lines that begin a type declaration, ignoring comment lines.
    pattern = re.compile(
        r"^(?:[ \t]*(?:public|internal|private|protected|sealed|abstract|"
        r"open|data|inline|value|annotation|fun|companion|actual|expect)\s+)*"
        r"(?:class|object|interface|enum\s+class)\s+(\w+)",
        re.MULTILINE,
    )
    m = pattern.search(text)
    return m.group(1) if m else kt_file.stem


def _package_label(kt_file: Path, src_root: Path) -> str:
    """Return a short package label based on subdirectory below src_root."""
    try:
        rel = kt_file.relative_to(src_root)
        parts = rel.parts
        if "mvt" in parts:
            return "mvt"
        if "ui" in parts:
            return "ui"
        return "main"
    except ValueError:
        return "main"


def _collect_components(repo_root: Path) -> list[ComponentInfo]:
    src_root = repo_root / "app" / "src" / "main" / "java"
    test_root = repo_root / "app" / "src" / "test" / "java"

    # Build a set of test file stems for fast lookup
    test_stems: set[str] = set()
    if test_root.exists():
        for t in test_root.rglob("*.kt"):
            # Convention: FooTest.kt tests Foo.kt
            stem = t.stem.removesuffix("Test")
            test_stems.add(stem)

    components: list[ComponentInfo] = []
    if not src_root.exists():
        return components

    for kt_file in sorted(src_root.rglob("*.kt")):
        class_name = _class_name_from_file(kt_file)
        pkg = _package_label(kt_file, src_root)
        has_tests = kt_file.stem in test_stems or class_name in test_stems
        rel = kt_file.relative_to(repo_root).as_posix()
        components.append(ComponentInfo(rel, class_name, pkg, has_tests))

    return components


def _collect_docs(repo_root: Path) -> list[Path]:
    docs_dir = repo_root / "docs"
    if not docs_dir.exists():
        return []
    return sorted(docs_dir.glob("*.md"))


# ---------------------------------------------------------------------------
# Markdown generation
# ---------------------------------------------------------------------------

_HEADER = """\
# CoreGuard Elite — Auto-Generated Software Requirements Document

**Document type:** Auto-generated SRD (repository introspection)  
**Application:** CoreGuard Elite  
**Generator:** `tools/generate_coreguard_srd.py`  
**Generated:** {date}  
**Status:** Informational — review before distributing externally

---
"""

_SECTION_OVERVIEW = """\
## 1. Product Overview

CoreGuard Elite is a native Kotlin Android security and device-monitoring
application. It provides:

- Runtime security checks (debugger, root, emulator, build type, signature).
- Nemesis forensic scanner (`mvt` package) with IOC matching against Pegasus
  and clandestine-spyware indicators.
- DNS sinkhole VPN service (`GuardVpnService`) for domain-level IOC blocking.
- Premium entitlement management via `BillingProvider` interface.
- Explainable Integrity Index (0–100 score) and evidence-based findings.

**Core philosophy:** Observe → Correlate → Explain → Respond → Verify  
**User-facing model:** Observe → Interpret → Protect → Verify

---
"""

_SECTION_BUILD = """\
## 2. Build Configuration (from `app/build.gradle.kts`)

| Field | Value |
|-------|-------|
| Namespace | `{namespace}` |
| Application ID | `{application_id}` |
| Min SDK | {min_sdk} |
| Compile SDK | {compile_sdk} |
| Target SDK | {target_sdk} |
| Version Code | {version_code} |
| Version Name | {version_name} |
| JVM Target | {jvm_target} |
| Release minification | {minify_release} |

---
"""


def _render_components_section(components: list[ComponentInfo]) -> str:
    groups: dict[str, list[ComponentInfo]] = {"main": [], "mvt": [], "ui": []}
    for c in components:
        groups.setdefault(c.package, []).append(c)

    lines: list[str] = ["## 3. Implemented Components\n"]
    label_map = {
        "main": "Core package (`com.coldboar.coreguard`)",
        "mvt": "MVT / Nemesis package (`com.coldboar.coreguard.mvt`)",
        "ui": "UI package (`com.coldboar.coreguard.ui`)",
    }
    for pkg_key in ("main", "mvt", "ui"):
        items = groups.get(pkg_key, [])
        if not items:
            continue
        lines.append(f"### {label_map.get(pkg_key, pkg_key)}\n")
        lines.append("| File | Primary class | Tests |")
        lines.append("|------|---------------|-------|")
        for c in items:
            tested = "✅" if c.has_tests else "—"
            lines.append(f"| `{c.relative_path}` | `{c.class_name}` | {tested} |")
        lines.append("")

    lines.append("---\n")
    return "\n".join(lines)


def _render_docs_section(docs: list[Path], repo_root: Path) -> str:
    if not docs:
        return ""
    lines = ["## 4. Reference Documentation\n"]
    lines.append("| Document | Path |")
    lines.append("|----------|------|")
    for d in docs:
        rel = d.relative_to(repo_root).as_posix()
        lines.append(f"| {d.stem.replace('_', ' ')} | `{rel}` |")
    lines.append("\n---\n")
    return "\n".join(lines)


_SECTION_CONSTRAINTS = """\
## 5. Implementation Constraints

1. **Preserve completed security logic.** Do not remove, rename, or weaken
   any evaluator, scanner, or VPN component without an explicit migration plan.
2. **Compile and test before broad new feature additions.** Run
   `./gradlew test assembleDebug` and confirm zero failures.
3. **Refactor safely rather than rewrite.** Prefer incremental changes.
4. **Distinguish implemented behavior from planned behavior.** Mark anything
   not yet validated in production with a clear status indicator.
5. **Never market heuristics as proof of malware or compromise.**
6. **Keep disruptive actions explicit, reversible, and user-controlled.**
7. **Do not rename** Quilla, Nemesis, Integrity Index, Integrity Timeline,
   Network Shield, or CoreGuard Elite without product-owner approval.

---
"""

_SECTION_PENDING = """\
## 6. Known Gaps / Pending Work

- Server-side billing entitlement verification (currently demo-only).
- APK signature hash not pinned in demo build (`SignatureCheckEvaluator`
  defaults `expectedSha256 = ""`).
- Full physical-device compatibility testing across SDK versions.
- Production IOC backend and key infrastructure.
- Final Play Store privacy policy and Data Safety declaration.
- Play Integrity API integration (anti-tampering).
- Backend JWT/token verification before production billing goes live.
"""


def generate_markdown(
    build: BuildConfig,
    components: list[ComponentInfo],
    docs: list[Path],
    repo_root: Path,
) -> str:
    date_str = datetime.date.today().isoformat()
    parts = [
        _HEADER.format(date=date_str),
        _SECTION_OVERVIEW,
        _SECTION_BUILD.format(
            namespace=build.namespace,
            application_id=build.application_id,
            min_sdk=build.min_sdk,
            compile_sdk=build.compile_sdk,
            target_sdk=build.target_sdk,
            version_code=build.version_code,
            version_name=build.version_name,
            jvm_target=build.jvm_target,
            minify_release="Enabled (R8/ProGuard)" if build.minify_release else "Disabled",
        ),
        _render_components_section(components),
        _render_docs_section(docs, repo_root),
        _SECTION_CONSTRAINTS,
        _SECTION_PENDING,
    ]
    return "".join(parts)


# ---------------------------------------------------------------------------
# Optional docx export (requires python-docx)
# ---------------------------------------------------------------------------

def write_docx(md_text: str, out_path: Path) -> None:
    try:
        from docx import Document  # type: ignore[import]
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required for --docx. Install it with: pip install python-docx"
        ) from exc

    document = Document()
    in_code = False

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            p = document.add_paragraph()
            p.add_run(line).font.name = "Courier New"
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            continue

        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            document.add_paragraph(text, style="List Bullet")
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            document.add_paragraph(text, style="List Number")
            continue

        if "|" in line and not line.strip().startswith("|---"):
            document.add_paragraph(line.replace("|", "  |  "))
            continue

        document.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate an up-to-date CoreGuard SRD from the repository."
    )
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=None,
        help="Repository root directory. Defaults to the parent of this script's directory.",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=Path("build/srd"),
        help="Output directory for generated artifacts (default: build/srd).",
    )
    parser.add_argument(
        "--docx",
        action="store_true",
        help="Also generate a .docx version of the SRD (requires python-docx).",
    )
    parser.add_argument(
        "--zip",
        action="store_true",
        help="Also bundle all generated artifacts into a .zip archive.",
    )
    args = parser.parse_args()

    # Resolve repo root relative to this script's location
    if args.repo_root is not None:
        repo_root = args.repo_root.resolve()
    else:
        repo_root = Path(__file__).resolve().parent.parent

    gradle_path = repo_root / "app" / "build.gradle.kts"
    if not gradle_path.exists():
        raise SystemExit(f"Cannot find app/build.gradle.kts under: {repo_root}")

    out_dir: Path = args.out_dir
    if not out_dir.is_absolute():
        out_dir = repo_root / out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    # Collect repository data
    build = _parse_build_config(gradle_path)
    components = _collect_components(repo_root)
    docs = _collect_docs(repo_root)

    # Generate markdown
    md_text = generate_markdown(build, components, docs, repo_root)
    md_out = out_dir / "CoreGuard_SRD.md"
    md_out.write_text(md_text, encoding="utf-8")
    print(f"Generated: {md_out}")

    generated: list[Path] = [md_out]

    # Optional docx
    if args.docx:
        docx_out = out_dir / "CoreGuard_SRD.docx"
        write_docx(md_text, docx_out)
        print(f"Generated: {docx_out}")
        generated.append(docx_out)

    # Optional zip
    if args.zip:
        zip_out = out_dir / "CoreGuard_SRD.zip"
        with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in generated:
                zf.write(f, arcname=f.name)
        print(f"Generated: {zip_out}")


if __name__ == "__main__":
    main()
